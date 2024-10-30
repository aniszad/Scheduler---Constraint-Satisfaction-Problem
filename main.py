
from constraint import Problem, MinConflictsSolver, BacktrackingSolver, RecursiveBacktrackingSolver
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor

schedule = {
    "sunday": 5,
    "monday": 5,
    "tuesday": 3,
    "wednesday": 5,
    "thursday": 5,
}

problem = Problem(MinConflictsSolver(steps=100000))

#problem = Problem(BacktrackingSolver) it takes a lot of time
#problem = Problem() it takes a lot of  time

# in this example we took 8 modules and 6 groups and two modules with a TP, which is the typical case,
# or you can use the input function in 'main'
modules = ["module1", "module2", "module3", "module4", "module5", "module6", "module7", "module8"]
modules_with_tp = ["module1", "module3"] # some module have a third session which is TP
groups = ["g1", "g2", "g3", "g4", "g5", "g6"]


variables = []

def define_variables():
    for module in modules:
        for group in groups:
            problem.addVariable(f"{module}_{group}_L", [f"{day}_{slot}" for day in schedule.keys() for slot in range(1, 3)])
            variables.append(f"{module}_{group}_L")
            problem.addVariable(f"{module}_{group}_T", [f"{day}_{slot}" for day in schedule.keys() for slot in range(3, schedule[day] + 1)])
            variables.append(f"{module}_{group}_T")
            if module in modules_with_tp:
                problem.addVariable(f"{module}_{group}_P", [f"{day}_{slot}" for day in schedule.keys() for slot in
                                                            range(3, schedule[day] + 1)])
                variables.append(f"{module}_{group}_P")


def diff_slot(*group_slots):
    return len(group_slots) == len(set(group_slots))


def enforce_diff_slots_for_every_group_constraint():
    by_group = {group: [] for group in groups}
    for variable in variables:
        module, group, session_type = variable.split("_")
        by_group[group].append(variable)
    for group, group_variables in by_group.items():
        problem.addConstraint(diff_slot, group_variables)

def sessions_within_two_days(*module_variables):
    days_used = set()

    for variable in module_variables:
        day, slot = variable.split("_")
        day = variable.split("_")[0]

        if day not in days_used:
            days_used.add(day)

        if len(days_used) > 2:
            return False

    return len(days_used) == 2

def enforce_sessions_within_two_days():
    by_module = {module: [] for module in modules}
    for variable in variables:
        module, group, session_type = variable.split("_")
        by_module[module].append(variable)
    for module, module_variables in by_module.items():
        problem.addConstraint(sessions_within_two_days, module_variables)

def same_lecture_slot(*lecture_variables):
    return len(set(lecture_variables)) == 1

def enforce_lecture_on_same_slot():
    by_lecture = {f"{module}": [] for module in modules}
    for variable in variables:
        if variable.endswith("_L"):
            by_lecture[variable.split('_')[0]].append(variable)
    print(by_lecture)
    for module_lecture, lecture_variables in by_lecture.items():
        problem.addConstraint(same_lecture_slot, lecture_variables)

def non_consecutive_slots(*group_variables):
    # Initialize a dictionary to group slots by day
    day_slots = {
        'sunday': [],
        'monday': [],
        'tuesday': [],
        'wednesday': [],
        'thursday': []
    }

    # Group slots by day
    for slot in group_variables:
        day, index = slot.split('_')
        day = day.lower()  # Convert day to lowercase to standardize

        if day in day_slots:
            day_slots[day].append(int(index))  # Append index as integer

    # Check each day for consecutive slots
    for day, slots in day_slots.items():
        if not slots:
            continue

        # Sort slots based on their index
        sorted_slots = sorted(slots)

        # Check for consecutive slots
        consecutive_count = 1  # At least one slot, so start count from 1
        max_consecutive = 1

        for i in range(1, len(sorted_slots)):
            if sorted_slots[i] == sorted_slots[i - 1] + 1:
                consecutive_count += 1
                max_consecutive = max(max_consecutive, consecutive_count)
            else:
                consecutive_count = 1

        if max_consecutive > 3:
            return False

    # If all checks pass, return True
    return True


def enforce_non_consecutive_slots():
    by_group = {group: [] for group in groups}
    for variable in variables:
        module, group, session_type = variable.split("_")
        by_group[group].append(variable)
    for group, group_variables in by_group.items():
        problem.addConstraint(non_consecutive_slots, group_variables)

def enforce_constraints():
    enforce_diff_slots_for_every_group_constraint()
    # enforce_sessions_within_two_days() #adding the constraint of two days per module also add a lot of time to the calculations
    enforce_lecture_on_same_slot()
    enforce_non_consecutive_slots()

def format_data(data):
    formatted_data = {
        "sunday": {
            f"g{i}": {f"slot{j}": "" for j in range(1, 6)} for i in range(1, 7)
        },
        "monday": {
            f"g{i}": {f"slot{j}": "" for j in range(1, 6)} for i in range(1, 7)
        },
        "tuesday": {
            f"g{i}": {f"slot{j}": "" for j in range(1, 6)} for i in range(1, 7)
        },
        "wednesday": {
            f"g{i}": {f"slot{j}": "" for j in range(1, 6)} for i in range(1, 7)
        },
        "thursday": {
            f"g{i}": {f"slot{j}": "" for j in range(1, 6)} for i in range(1, 7)
        }
    }

    for key, value in data.items():
        module, group, info_type = key.split('_')
        day, slot = value.split('_')

        slot_number = f"slot{slot}"
        module_type = f"{module}_{info_type.capitalize()}"

        formatted_data[day][f"g{group[1]}"][slot_number] = module_type

    return formatted_data

def build_day_table(day_data):
    table = [['' for _ in range(6)] for _ in range(6)]

    for group, slots_data in day_data.items():
        group_index = int(group[1]) - 1
        table[group_index][0] = f"Group {group[1]}"

        for slot, module_type in slots_data.items():
            slot_index = int(slot[-1])

            if module_type:
                table[group_index][slot_index] = module_type

    return table

def generate_day_tables(data):
    day_tables = {}

    for day, day_data in data.items():
        table = build_day_table(day_data)
        day_tables[day] = table

    return day_tables


def add_table_to_docx(doc, day, table):
    # Define a color map for modules
    module_colors = {
        "module1": "FF0000",  # Red
        "module2": "00FF00",  # Green
        "module3": "0000FF",  # Blue
        "module4": "FFFF00",  # Yellow
        "module5": "FF00FF",  # Magenta
        "module6": "00FFFF",  # Cyan
        "module7": "800080",  # Purple
        "module8": "808000"  # Olive
    }

    doc.add_heading(f"{day.capitalize()} Table", level=1)
    table_rows = len(table)
    table_cols = len(table[0])
    doc_table = doc.add_table(rows=table_rows, cols=table_cols)

    for i in range(table_rows):
        row = doc_table.rows[i]
        for j in range(table_cols):
            cell = row.cells[j]
            cell.text = table[i][j]

            # Apply color to the cell based on the module
            if cell.text:
                module_name = cell.text.split('_')[0]
                if module_name in module_colors:
                    color = module_colors[module_name]
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                    cell._element.get_or_add_tcPr().append(shading_elm)

def create_docx_with_tables(day_tables, output_filename):
    doc = Document()

    for day, table in day_tables.items():
        add_table_to_docx(doc, day, table)
        doc.add_paragraph("")

    doc.save(output_filename)


def get_user_input():
    # Input modules
    modules = []
    while True:
        module_input = input("Enter module name (enter 'done' when finished): ").strip()
        if module_input.lower() == 'done':
            break
        modules.append(module_input)

    # Input modules with TP
    modules_with_tp = []
    tp_input = input("Enter modules with third session (TP) (comma-separated, enter 'none' if none): ").strip()
    if tp_input.lower() != 'none':
        modules_with_tp = [module.strip() for module in tp_input.split(',')]

    # Input groups
    groups = []
    while len(groups) < 6:
        group_input = input(f"Enter group name (at least {len(groups) + 1} groups required): ").strip()
        if group_input not in groups:
            groups.append(group_input)

    return modules, modules_with_tp, groups

if __name__ == '__main__':
    #modules, modules_with_tp, groups = get_user_input() we can use the suer input instead of the example
    define_variables()
    enforce_constraints()
    solution = problem.getSolution()
    if solution:
        print(solution)
        formatted_schedule = format_data(solution)
        print(formatted_schedule)
        day_tables = generate_day_tables(formatted_schedule)
        create_docx_with_tables(day_tables, "output.docx")
    else:
        print("No solution found")
